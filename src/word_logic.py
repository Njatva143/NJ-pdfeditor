from docx import Document
import io

def process_word_file(file_stream, search_text, replace_text):
    try:
        # Load the document
        doc = Document(file_stream)
        replaced_count = 0

        # 1. Paragraphs me dhundho
        for paragraph in doc.paragraphs:
            if search_text in paragraph.text:
                # Text replace karo (Run level replacement is hard, doing paragraph level)
                # Note: Complex formatting might reset in the specific line being changed
                paragraph.text = paragraph.text.replace(search_text, replace_text)
                replaced_count += 1

        # 2. Tables me bhi dhundho
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if search_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(search_text, replace_text)
                            replaced_count += 1
        
        # Save to buffer
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        return output_buffer.getvalue(), replaced_count
        
    except Exception as e:
        return None, str(e)
